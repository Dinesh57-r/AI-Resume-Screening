"""
parser.py – Resume file parser for PDF and DOCX formats.
"""
import io
import re


def parse_pdf(file_bytes: bytes) -> str:
    """Extract raw text from a PDF file's bytes using pdfplumber + pypdf fallback."""
    text_plumber = ""
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text() or page.extract_text(layout=True)
                if page_text:
                    text_parts.append(page_text)
        text_plumber = "\n".join(text_parts).strip()
    except Exception:
        text_plumber = ""

    # If pdfplumber extracted sufficient text (> 50 chars), return it
    if len(text_plumber) > 50:
        return text_plumber

    # Fallback to pypdf engine
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(file_bytes))
        pypdf_parts = []
        for page in reader.pages:
            t = page.extract_text()
            if t:
                pypdf_parts.append(t)
        text_pypdf = "\n".join(pypdf_parts).strip()
        if len(text_pypdf) > len(text_plumber):
            return text_pypdf
    except Exception:
        pass

    return text_plumber if text_plumber else "[PDF parse error: Could not extract text from file]"


def parse_docx(file_bytes: bytes) -> str:
    """Extract raw text from a DOCX file's bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[DOCX parse error: {e}]"


def parse_txt(file_bytes: bytes) -> str:
    """Decode plain text file."""
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        return f"[TXT parse error: {e}]"


def parse_resume(uploaded_file) -> str:
    """
    Dispatch parser based on file extension.
    
    Args:
        uploaded_file: Streamlit UploadedFile object.
    Returns:
        Extracted raw text string.
    """
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()

    if name.endswith(".pdf"):
        return parse_pdf(file_bytes)
    elif name.endswith(".docx") or name.endswith(".doc"):
        return parse_docx(file_bytes)
    elif name.endswith(".txt"):
        return parse_txt(file_bytes)
    else:
        # Try PDF first, then DOCX
        text = parse_pdf(file_bytes)
        if "[PDF parse error" not in text and len(text) > 50:
            return text
        return parse_txt(file_bytes)


def clean_text(text: str) -> str:
    """Remove excessive whitespace and normalize line endings."""
    if not text:
        return ""
    # Normalize whitespace
    text = re.sub(r"\r\n|\r", "\n", text)
    # Remove multiple blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)
    # Remove excessive spaces
    text = re.sub(r" {2,}", " ", text)
    return text.strip()
